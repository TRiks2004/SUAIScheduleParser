from docx import Document
from docx.table import Table

from typing import Literal, List

from model import *
from parser_t import Persons


def get_name(text: str) -> tuple[str, str] | None:
    if text.count('-') > 2:
        return None
    else:
        bloks = text.split()

        for i, blok in enumerate(bloks):
            if blok.count('.') == 2 and len(blok) == 4:
                return ' '.join(bloks[:i-1]), ' '.join(bloks[i-1:])  

def get_full_date_text(text: str, classroom: Classroom) -> str:
    persons = Persons()
    
    return persons.get_teachers(text, classroom)

def get_teachers(text: str) -> List[Teacher]:

    bloks = text.split()

    teachers_str: List[str] = []

    k: int = 0

    for i, blok in enumerate(bloks):
        if blok.count('.') == 2 and len(blok) == 4:
            teachers_str.append(' '.join([bloks[i-1], bloks[i]]))
            k = i

    aud_num = bloks[k+1].split('.')[1].split(',')

    point = list(zip(teachers_str, aud_num))

    return [get_full_date_text(text=i[0], classroom=Classroom(name=i[1])) for i in point]
        

def get_class(number: int, text: str) -> Class:

    name_dop = get_name(text)

    if name_dop is None:
        return Class(number=number, name=None, teachers=None, classroom=None)
    else:
        name, dop = name_dop

        teachers = get_teachers(dop)
        print(f'{dop=}\n{teachers=}\n\n\n\n')    
        return Class(number=number, name=name, teachers=teachers)

def get_schedule(table: Table):
    # .rows[0].cells[0].text

    day_list: List[Day] = []
    

    for i, col in enumerate(table.columns):
        
        if i in [0, 1]: continue
        
        class_list: List[Class] = []
        for number, cell in enumerate(col.cells[1:]):            
            class_list.append(get_class(number + 1, cell.text))

        day_list.append(
            Day(week_day=col.cells[0].text, classes=class_list)
        )
    
    return Schedule(days=day_list)


def get_schedule_group(all_full):
    return [ScheduleGroup(get_hender(table[0].text), get_schedule(table[1])) for table in all_full]


def main():
    path = 'sh.docx'

    doc = Document(path)

    all_paragraphs = doc.paragraphs
    all_tables = doc.tables

    all_full = list(zip(all_paragraphs[2:], all_tables))

    full_schedule = get_schedule_group(all_full)



    for schedule_group in full_schedule:
        hender_text = [f'{"":>25}']

        row = [[f'{1:>25}', ], [f'{2:>25}', ], [f'{3:>25}', ], [f'{4:>25}', ]]


        for schedule in schedule_group.schedule.days:
            hender_text.append(f'{schedule.week_day:>25}')

            for clas in schedule.classes:
                if clas.teachers is None: 
                    row[clas.number-1].append(f'{"----":>25}')
                    continue

                row[clas.number-1].append(f'{str(clas.name[:19]):>25}')
                
            

        h = '|'.join(hender_text)
        print(f'{schedule_group.hender}')
        print()
        print(h)
        
        print('—'*len(h))
        for i in row:
            print('|'.join(i))
            print('—'*len(h))

        print()
main()